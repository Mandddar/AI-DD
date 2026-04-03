"""
Text extraction from uploaded documents.
Supports: PDF, Word (.docx), Excel (.xlsx/.xls), CSV, TSV, plain text, Images.

Libraries:
  - pymupdf (fitz) — PDF parsing + PDF-to-image for OCR fallback
  - python-docx — Word .docx parsing
  - openpyxl — Excel .xlsx parsing
  - pandas — CSV / TSV parsing
  - pytesseract + Pillow — OCR for scanned/image documents
  - google-cloud-documentai — Google Document AI (primary OCR, spec §2.4)
"""
import io
import logging
import os
from pathlib import Path

logger = logging.getLogger(__name__)

# Google Document AI configuration (spec §2.4)
_GOOGLE_DOC_AI_ENABLED = bool(os.getenv("GOOGLE_DOCUMENT_AI_PROCESSOR_ID"))
_GOOGLE_PROJECT_ID = os.getenv("GOOGLE_CLOUD_PROJECT", "")
_GOOGLE_LOCATION = os.getenv("GOOGLE_DOCUMENT_AI_LOCATION", "eu")  # EU region per spec §2.3
_GOOGLE_PROCESSOR_ID = os.getenv("GOOGLE_DOCUMENT_AI_PROCESSOR_ID", "")


def extract_text(file_bytes: bytes, mime_type: str, filename: str) -> tuple[str, str | None]:
    """
    Returns (extracted_text, page_count_str).
    page_count_str is None for non-paginated formats.
    """
    ext = Path(filename).suffix.lower()

    if mime_type == "application/pdf" or ext == ".pdf":
        return _extract_pdf(file_bytes)

    if mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ) or ext in (".docx", ".doc"):
        return _extract_docx(file_bytes), None

    if mime_type in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    ) or ext in (".xlsx", ".xls"):
        return _extract_excel(file_bytes), None

    if ext in (".csv", ".tsv") or mime_type in ("text/csv", "text/tab-separated-values"):
        return _extract_csv_tsv(file_bytes, ext), None

    if mime_type.startswith("text/") or ext == ".txt":
        return file_bytes.decode("utf-8", errors="replace"), None

    if mime_type.startswith("image/") or ext in (".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp"):
        return _extract_image_ocr(file_bytes, mime_type), None

    return "", None


def _extract_pdf(data: bytes) -> tuple[str, str]:
    """Extract text from PDF using pymupdf. Falls back to OCR for scanned pages."""
    import fitz  # pymupdf

    doc = fitz.open(stream=data, filetype="pdf")
    page_count = doc.page_count
    pages = []
    for page in doc:
        text = page.get_text()
        if text.strip():
            pages.append(text)
        else:
            # Scanned page — try Google Document AI first, then pytesseract fallback
            ocr_text = _ocr_scanned_page(page, data)
            if ocr_text.strip():
                pages.append(ocr_text)
    doc.close()
    return "\n\n".join(pages), str(page_count)


def _ocr_scanned_page(page, full_pdf_bytes: bytes) -> str:
    """OCR a scanned page — Google Document AI primary, pytesseract fallback."""
    # Try Google Document AI first if configured
    if _GOOGLE_DOC_AI_ENABLED:
        try:
            result = _google_document_ai_ocr(full_pdf_bytes, "application/pdf")
            if result.strip():
                return result
        except Exception as e:
            logger.warning("Google Document AI failed, falling back to Tesseract: %s", e)

    # Pytesseract fallback
    try:
        pix = page.get_pixmap(dpi=300)
        img_bytes = pix.tobytes("png")
        return _ocr_tesseract(img_bytes)
    except Exception as e:
        logger.warning("Tesseract OCR fallback failed for page %d: %s", page.number, e)
    return ""


def _extract_docx(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _extract_excel(data: bytes) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            row_text = "\t".join(str(c) if c is not None else "" for c in row)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def _extract_csv_tsv(data: bytes, ext: str) -> str:
    """Parse CSV/TSV files using pandas."""
    import pandas as pd
    sep = "\t" if ext == ".tsv" else ","
    try:
        df = pd.read_csv(io.BytesIO(data), sep=sep)
        return df.to_string(index=False)
    except Exception as e:
        logger.warning("pandas CSV/TSV parse failed: %s, falling back to raw text", e)
        return data.decode("utf-8", errors="replace")


def _extract_image_ocr(data: bytes, mime_type: str = "image/png") -> str:
    """OCR an image file — Google Document AI primary, pytesseract fallback."""
    if _GOOGLE_DOC_AI_ENABLED:
        try:
            result = _google_document_ai_ocr(data, mime_type)
            if result.strip():
                return result
        except Exception as e:
            logger.warning("Google Document AI failed for image, falling back to Tesseract: %s", e)

    return _ocr_tesseract(data)


# ── Google Document AI (primary OCR per spec §2.4) ────────

def _google_document_ai_ocr(file_bytes: bytes, mime_type: str) -> str:
    """
    Process a document using Google Document AI (spec §2.4).
    Requires: GOOGLE_DOCUMENT_AI_PROCESSOR_ID, GOOGLE_CLOUD_PROJECT env vars.
    Uses EU region per spec §2.3 (europe-west3).
    """
    from google.cloud import documentai_v1 as documentai
    from google.api_core.client_options import ClientOptions

    opts = ClientOptions(api_endpoint=f"{_GOOGLE_LOCATION}-documentai.googleapis.com")
    client = documentai.DocumentProcessorServiceClient(client_options=opts)

    resource_name = client.processor_path(_GOOGLE_PROJECT_ID, _GOOGLE_LOCATION, _GOOGLE_PROCESSOR_ID)

    raw_document = documentai.RawDocument(content=file_bytes, mime_type=mime_type)
    request = documentai.ProcessRequest(name=resource_name, raw_document=raw_document)

    result = client.process_document(request=request)
    return result.document.text


# ── Pytesseract (fallback OCR) ────────────────────────────

def _ocr_tesseract(img_bytes: bytes) -> str:
    """Run pytesseract OCR on raw image bytes (fallback for when Google Document AI is unavailable)."""
    try:
        from PIL import Image
        import pytesseract
        image = Image.open(io.BytesIO(img_bytes))
        return pytesseract.image_to_string(image, lang="eng")
    except Exception as e:
        logger.error("pytesseract OCR failed: %s", e)
        return ""
