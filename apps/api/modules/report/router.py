"""
Report Module — API endpoints for generating and managing DD reports.

Tech: FastAPI + python-docx + openpyxl + jinja2 + Groq API
Report types: Detailed Workstream, Executive Summary, Consolidated Overall
"""
from uuid import UUID
from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import FileResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from core.database import get_db
from modules.auth.dependencies import project_manager, project_contributor, project_reader
from modules.auth.models import User, UserRole
from .models import Report, ReportType, ReportFormat
from .schemas import ReportCreate, ReportContentUpdate, ReportOut

router = APIRouter(prefix="/projects/{project_id}/reports", tags=["reports"])


@router.post("/generate", response_model=ReportOut, status_code=status.HTTP_201_CREATED)
async def generate_report(
    project_id: UUID,
    data: ReportCreate,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(project_manager),
):
    """Generate a new report populated with AI findings from completed analyses."""
    content = await _build_report_content(project_id, data.report_type, data.workstream, db)

    # If no completed analysis or no findings, return error instead of empty report
    if "Notice" in content and len(content) == 1:
        raise HTTPException(status_code=400, detail=content["Notice"])

    report = Report(
        project_id=project_id,
        created_by=user.id,
        report_type=data.report_type,
        report_format=data.report_format if data.report_format in ("docx", "xlsx") else "docx",
        workstream=data.workstream,
        title=data.title,
        content=content,
    )
    db.add(report)
    await db.commit()
    await db.refresh(report)
    return report


async def _build_report_content(
    project_id: UUID,
    report_type: str,
    workstream: str | None,
    db: AsyncSession,
) -> dict:
    """Pull AI findings from agent_findings across ALL completed runs for this project."""
    from modules.agent.models import AgentFinding, AgentRun, RunStatus

    # Get ALL completed runs for this project (not just the latest)
    runs_result = await db.execute(
        select(AgentRun.id)
        .where(AgentRun.project_id == project_id)
        .where(AgentRun.status == RunStatus.completed)
        .order_by(AgentRun.completed_at.desc())
    )
    completed_run_ids = [row[0] for row in runs_result.all()]

    if not completed_run_ids:
        return {"Notice": "No completed AI analysis found. Run AI analysis on this project first, then try generating the report."}

    # Get findings from ALL completed runs — filter by workstream if detailed report
    findings_query = select(AgentFinding).where(AgentFinding.run_id.in_(completed_run_ids))
    if report_type == "detailed_workstream" and workstream:
        findings_query = findings_query.where(AgentFinding.agent_type == workstream)
    findings_query = findings_query.order_by(AgentFinding.severity.desc())

    findings_result = await db.execute(findings_query)
    all_findings = list(findings_result.scalars().all())

    # Deduplicate by title (keep earliest occurrence = from latest run since ordered desc)
    seen_titles: set[str] = set()
    findings = []
    for f in all_findings:
        if f.title not in seen_titles:
            seen_titles.add(f.title)
            findings.append(f)

    if not findings:
        ws_name = (workstream or "").capitalize()
        if report_type == "detailed_workstream" and workstream:
            return {"Notice": f"No {ws_name} findings available. Please run AI analysis on the '{ws_name}' workstream first, then try generating this report again."}
        return {"Notice": "No findings found for this project. Run AI analysis first, then generate the report."}

    # Build structured content
    severity_order = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}

    if report_type == "executive_summary":
        # Condensed summary — critical and high findings only
        key_findings = [f for f in findings if f.severity in ("critical", "high")]
        content = {
            "Executive Summary": f"This report covers {len(findings)} findings from the AI-assisted due diligence analysis. "
                                 f"{len(key_findings)} findings are classified as critical or high severity.",
            "Critical & High Severity Findings": [
                {
                    "Title": f.title,
                    "Category": f.category,
                    "Severity": f.severity.upper(),
                    "Workstream": f.agent_type,
                    "Description": f.description,
                }
                for f in key_findings
            ] if key_findings else "No critical or high severity findings identified.",
            "Finding Summary by Severity": {
                severity: len([f for f in findings if f.severity == severity])
                for severity in ["critical", "high", "medium", "low", "info"]
                if any(f.severity == severity for f in findings)
            },
        }
    elif report_type == "consolidated":
        # Group by workstream
        workstreams_found = sorted(set(f.agent_type for f in findings))
        content = {
            "Overall Assessment": f"Consolidated report covering {len(findings)} findings across {len(workstreams_found)} workstream(s): {', '.join(workstreams_found)}.",
        }
        for ws in workstreams_found:
            ws_findings = [f for f in findings if f.agent_type == ws]
            ws_findings.sort(key=lambda f: severity_order.get(f.severity, 99))
            content[f"{ws.capitalize()} Workstream ({len(ws_findings)} findings)"] = [
                {
                    "Title": f.title,
                    "Severity": f.severity.upper(),
                    "Category": f.category,
                    "Description": f.description,
                }
                for f in ws_findings
            ]
    else:
        # Detailed workstream report
        ws_label = (workstream or "all").capitalize()
        # Group by category
        categories = {}
        for f in findings:
            cat = f.category or "General"
            if cat not in categories:
                categories[cat] = []
            categories[cat].append(f)

        content = {
            f"{ws_label} Due Diligence Report": f"Detailed findings for the {ws_label} workstream. Total: {len(findings)} findings.",
        }
        for cat, cat_findings in categories.items():
            cat_findings.sort(key=lambda f: severity_order.get(f.severity, 99))
            content[cat] = [
                {
                    "Title": f.title,
                    "Severity": f.severity.upper(),
                    "Description": f.description,
                    "Source": f.source_excerpts[0] if f.source_excerpts else "—",
                }
                for f in cat_findings
            ]

    content["AI Disclaimer"] = (
        "This report was generated using Artificial Intelligence. AI-generated results may be inaccurate, "
        "incomplete, or misleading. Responsibility for audit results, their interpretation, and all decisions "
        "derived therefrom lies exclusively with the human reviewer. This tool does not replace qualified "
        "legal, tax, or financial advisory services."
    )

    return content


@router.get("/", response_model=list[ReportOut])
async def list_reports(
    project_id: UUID,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(project_reader),
):
    """List all reports for a project."""
    result = await db.execute(
        select(Report)
        .where(Report.project_id == project_id)
        .order_by(Report.created_at.desc())
    )
    return list(result.scalars().all())


@router.get("/{report_id}", response_model=ReportOut)
async def get_report(
    project_id: UUID,
    report_id: UUID,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(project_reader),
):
    """Get a specific report with its content."""
    report = await db.get(Report, report_id)
    if not report or report.project_id != project_id:
        raise HTTPException(status_code=404, detail="Report not found")
    if user.role == UserRole.buyer and not report.is_finalized:
        raise HTTPException(status_code=403, detail="Buyers can only access finalized reports")
    return report


@router.patch("/{report_id}/edit", response_model=ReportOut)
async def edit_report_content(
    project_id: UUID,
    report_id: UUID,
    update: ReportContentUpdate,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(project_manager),
):
    """Edit report content (executive summary editing before export)."""
    report = await db.get(Report, report_id)
    if not report or report.project_id != project_id:
        raise HTTPException(status_code=404, detail="Report not found")
    if report.is_finalized:
        raise HTTPException(status_code=400, detail="Cannot edit finalized report")

    report.edited_content = update.edited_content
    await db.commit()
    await db.refresh(report)
    return report


@router.post("/{report_id}/finalize", response_model=ReportOut)
async def finalize_report(
    project_id: UUID,
    report_id: UUID,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(project_manager),
):
    """Finalize report — generates the .docx file for download."""
    report = await db.get(Report, report_id)
    if not report or report.project_id != project_id:
        raise HTTPException(status_code=404, detail="Report not found")

    # Generate file in requested format
    from datetime import datetime, timezone
    if report.report_format == ReportFormat.xlsx:
        file_path = _generate_xlsx(report)
    else:
        file_path = _generate_docx(report)
    report.storage_path = file_path
    report.is_finalized = True
    report.finalized_by = user.id
    report.finalized_at = datetime.now(timezone.utc)
    await db.commit()
    await db.refresh(report)
    return report


@router.get("/{report_id}/download")
async def download_report(
    project_id: UUID,
    report_id: UUID,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(project_reader),
):
    """Download finalized report as .docx."""
    report = await db.get(Report, report_id)
    if not report or report.project_id != project_id:
        raise HTTPException(status_code=404, detail="Report not found")
    if user.role == UserRole.buyer and not report.is_finalized:
        raise HTTPException(status_code=403, detail="Buyers can only access finalized reports")
    if not report.storage_path:
        raise HTTPException(status_code=400, detail="Report file not yet generated. Finalize the report first.")

    import os
    if not os.path.exists(report.storage_path):
        raise HTTPException(status_code=404, detail="Report file not found on disk")

    safe_title = "".join(c if c.isalnum() or c in " -_" else "_" for c in report.title)

    if report.report_format == ReportFormat.xlsx:
        return FileResponse(
            path=report.storage_path,
            filename=f"{safe_title}.xlsx",
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

    return FileResponse(
        path=report.storage_path,
        filename=f"{safe_title}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def _generate_docx(report) -> str:
    """Generate a Word .docx file from report content."""
    from docx import Document as DocxDocument
    from pathlib import Path
    import uuid

    doc = DocxDocument()

    # Title
    doc.add_heading(report.title, level=0)

    # AI Disclaimer
    disclaimer = doc.add_paragraph()
    disclaimer.add_run(
        "Notice: This report uses Artificial Intelligence to support the due diligence review. "
        "AI-generated results may be inaccurate, incomplete, or misleading. Responsibility for "
        "audit results lies exclusively with the human reviewer."
    ).italic = True

    doc.add_paragraph("")

    # Report metadata
    from datetime import datetime, timezone
    meta = doc.add_paragraph()
    meta.add_run(f"Report Type: ").bold = True
    meta.add_run(f"{report.report_type}\n")
    if report.workstream:
        meta.add_run(f"Workstream: ").bold = True
        meta.add_run(f"{report.workstream}\n")
    meta.add_run(f"Generated: ").bold = True
    meta.add_run(f"{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}\n")

    doc.add_paragraph("")

    # Content — use edited_content if available, otherwise content
    content = report.edited_content or report.content or {}
    if isinstance(content, dict):
        for section_title, section_body in content.items():
            doc.add_heading(section_title, level=1)
            if isinstance(section_body, str):
                doc.add_paragraph(section_body)
            elif isinstance(section_body, list):
                for item in section_body:
                    if isinstance(item, dict):
                        for k, v in item.items():
                            p = doc.add_paragraph()
                            p.add_run(f"{k}: ").bold = True
                            p.add_run(str(v))
                    else:
                        doc.add_paragraph(str(item), style="List Bullet")
            elif isinstance(section_body, dict):
                for k, v in section_body.items():
                    p = doc.add_paragraph()
                    p.add_run(f"{k}: ").bold = True
                    p.add_run(str(v))
            else:
                doc.add_paragraph(str(section_body))
    else:
        doc.add_paragraph(str(content))

    # Save
    uploads_dir = Path(__file__).parent.parent.parent / "uploads" / "reports"
    uploads_dir.mkdir(parents=True, exist_ok=True)
    filepath = uploads_dir / f"{uuid.uuid4()}.docx"
    doc.save(str(filepath))
    return str(filepath)


def _generate_xlsx(report) -> str:
    """Generate an Excel .xlsx file from report content (spec §12.2)."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from pathlib import Path
    import uuid

    wb = Workbook()

    # Remove default sheet
    wb.remove(wb.active)

    content = report.edited_content or report.content or {}

    # Styles
    header_font = Font(bold=True, size=12, color="FFFFFF")
    header_fill = PatternFill(start_color="1a1a2e", end_color="1a1a2e", fill_type="solid")
    gold_font = Font(bold=True, size=11, color="C9A84C")
    thin_border = Border(
        bottom=Side(style="thin", color="CCCCCC"),
    )

    if isinstance(content, dict):
        # Group sections — create one sheet per major section or workstream
        section_items = list(content.items())

        # If findings are grouped by workstream, create per-workstream sheets
        # Otherwise create a single "Report" sheet
        for section_title, section_body in section_items:
            # Sanitize sheet name (max 31 chars, no special chars)
            sheet_name = section_title[:31].replace("/", "-").replace("\\", "-").replace("*", "").replace("?", "").replace("[", "").replace("]", "")
            if not sheet_name:
                sheet_name = "Report"

            ws = wb.create_sheet(title=sheet_name)

            # Title row
            ws.append([report.title])
            ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=5)
            ws.cell(1, 1).font = Font(bold=True, size=14, color="C9A84C")

            # Section header
            ws.append([section_title])
            ws.cell(2, 1).font = gold_font
            ws.append([])  # blank row

            if isinstance(section_body, str):
                ws.append([section_body])
            elif isinstance(section_body, list):
                # Findings list — create table
                if section_body and isinstance(section_body[0], dict):
                    keys = list(section_body[0].keys())
                    # Header row
                    row_num = ws.max_row + 1
                    for col_idx, key in enumerate(keys, 1):
                        cell = ws.cell(row=row_num, column=col_idx, value=key)
                        cell.font = header_font
                        cell.fill = header_fill
                        cell.alignment = Alignment(horizontal="center")

                    # Data rows
                    for item in section_body:
                        row_num = ws.max_row + 1
                        for col_idx, key in enumerate(keys, 1):
                            cell = ws.cell(row=row_num, column=col_idx, value=str(item.get(key, "")))
                            cell.border = thin_border
                else:
                    for item in section_body:
                        ws.append([str(item)])
            elif isinstance(section_body, dict):
                for k, v in section_body.items():
                    ws.append([str(k), str(v)])
            else:
                ws.append([str(section_body)])

            # Auto-width columns
            for col in ws.columns:
                max_len = 0
                col_letter = col[0].column_letter
                for cell in col:
                    if cell.value:
                        max_len = max(max_len, len(str(cell.value)))
                ws.column_dimensions[col_letter].width = min(max_len + 4, 60)

    if len(wb.sheetnames) == 0:
        ws = wb.create_sheet("Report")
        ws.append([str(content)])

    # Save
    uploads_dir = Path(__file__).parent.parent.parent / "uploads" / "reports"
    uploads_dir.mkdir(parents=True, exist_ok=True)
    filepath = uploads_dir / f"{uuid.uuid4()}.xlsx"
    wb.save(str(filepath))
    return str(filepath)
